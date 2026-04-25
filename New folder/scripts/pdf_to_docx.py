from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from pypdf import PdfReader


HEADING_RE = re.compile(r"^\d+(?:\.\d+)*\.?\s+[A-Z][A-Z0-9 ,&()/-]+$")


def add_page_content(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split())
        if not line:
            continue

        if line.isupper() and len(line) <= 120:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            if HEADING_RE.match(line):
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(11)
            continue

        paragraph = document.add_paragraph(line)
        paragraph.style = document.styles["Normal"]


def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> None:
    reader = PdfReader(str(pdf_path))
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        add_page_content(document, text)

        if index < len(reader.pages) - 1:
            document.add_section(WD_SECTION.NEW_PAGE)

    document.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: pdf_to_docx.py <input.pdf> <output.docx>")
        return 1

    pdf_path = Path(sys.argv[1]).expanduser().resolve()
    output_path = Path(sys.argv[2]).expanduser().resolve()

    if not pdf_path.exists():
        print(f"Input PDF not found: {pdf_path}")
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    convert_pdf_to_docx(pdf_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
